"""
文档解析:PDF / Word / Excel / txt → 结构化文本块(带来源定位)

面试点:扫描件 PDF 无文字层时需 OCR(paddleocr/rapidocr)。
当前版本对无文字页直接跳过,OCR 作为扩展点不做死。
"""
import os
from dataclasses import dataclass


@dataclass
class ParsedBlock:
    source: str   # 文件名(绝对路径,引用时转相对路径)
    section: str  # 来源定位:页码 / 工作表 / 标题,用于引用
    text: str


def parse_document(path: str) -> list[ParsedBlock]:
    """按扩展名分发解析,返回带来源定位的文本块"""
    ext = os.path.splitext(path)[1].lower()
    parser = {
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".xlsx": _parse_xlsx,
        ".txt": _parse_txt,
    }.get(ext)
    if parser is None:
        print(f"⚠️ 跳过不支持的类型: {path}")
        return []
    return parser(path)


def _parse_pdf(path: str) -> list[ParsedBlock]:
    from pypdf import PdfReader
    reader = PdfReader(path)
    blocks = []
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if len(text.strip()) >= 10:  # 过滤空白/纯图片页
            blocks.append(ParsedBlock(path, f"第{i}页", text))
    return blocks


def _parse_docx(path: str) -> list[ParsedBlock]:
    import docx
    doc = docx.Document(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    for table in doc.tables:  # 表格也并入,保证信息不丢
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            text += "\n" + " | ".join(cells)
    return [ParsedBlock(path, "正文", text)] if text.strip() else []


def _parse_xlsx(path: str) -> list[ParsedBlock]:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    blocks = []
    for ws in wb.worksheets:
        lines = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append(" | ".join(cells))
        if lines:
            blocks.append(ParsedBlock(path, f"工作表[{ws.title}]", "\n".join(lines)))
    wb.close()
    return blocks


def _parse_txt(path: str) -> list[ParsedBlock]:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [ParsedBlock(path, "全文", text)] if text.strip() else []
